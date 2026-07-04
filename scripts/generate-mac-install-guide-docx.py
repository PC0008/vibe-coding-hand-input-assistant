#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Vibe Coding手持输入助手-Mac安装与常见提示处理.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Vibe Coding手持输入助手")
    style_run(run, bold=True, color="0B2545", size=24)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Mac 安装与常见安全提示处理指南")
    style_run(run, color="555555", size=13)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.add_run("适用对象：").bold = True
    p.add_run("朋友测试版用户、早期体验用户、需要安装 DMG 并连接 StickS3 手持硬件的 Mac 用户。")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string("2E74B5" if level <= 2 else "1F4D78")
    return p


def add_note(doc, title, body, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.4)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    style_run(r, bold=True, color="1F3A5F", size=11)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_steps(doc, steps):
    for idx, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(step)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_command(doc, command):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(command)
    r.font.name = "Menlo"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
    r.font.size = Pt(9.5)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25

    add_title(doc)
    add_note(
        doc,
        "先看结论",
        "当前测试版没有 Apple Developer ID 公证。首次从浏览器、微信、网盘下载后，macOS 可能提示“无法验证开发者”或“已损坏，无法打开”。这通常不是软件真的坏了，而是 Gatekeeper 安全机制拦截。"
    )

    add_heading(doc, "一、标准安装流程")
    add_steps(doc, [
        "下载 VibeCodingHandInputAssistant-0.1.0.dmg。",
        "双击打开 DMG。",
        "把“Vibe Coding手持输入助手”拖到 Applications / 应用程序。",
        "到“应用程序”里右键点击 App，选择“打开”。不要第一次直接双击。",
        "如果系统提示来自未验证开发者，选择“打开”或到“隐私与安全性”里点“仍要打开”。",
        "进入系统设置 -> 隐私与安全性 -> 辅助功能，打开“Vibe Coding手持输入助手”。",
        "打开 App 后，选择目标软件，连接蓝牙设备 Vibe Coding Remote。"
    ])

    add_heading(doc, "二、如果提示“已损坏，无法打开”")
    doc.add_paragraph("如果看到类似提示：")
    add_note(doc, "系统提示示例", "“Vibe Coding手持输入助手”已损坏，无法打开。你应该将它移到废纸篓。", fill="FFF4E5")
    doc.add_paragraph("请按下面步骤处理：")
    add_steps(doc, [
        "确认 App 已经拖到“应用程序”。",
        "打开“终端”。",
        "复制并执行下面的命令。",
        "回到“应用程序”，右键 App，选择“打开”。"
    ])
    add_command(doc, 'xattr -dr com.apple.quarantine "/Applications/Vibe Coding手持输入助手.app"')
    doc.add_paragraph("如果仍然打不开，可以继续在终端执行：")
    add_command(doc, 'open "/Applications/Vibe Coding手持输入助手.app"')

    add_heading(doc, "三、如果辅助功能显示未开启")
    add_bullets(doc, [
        "先确认你授权的是 /Applications 里的 App，不是 DMG 里临时打开的 App。",
        "如果系统设置里开关是蓝色，但 App 里仍显示未开启，先点减号删除旧条目。",
        "再点加号，重新选择 /Applications/Vibe Coding手持输入助手.app。",
        "重新打开 App。测试版使用 ad-hoc 签名，覆盖安装新版本后可能需要重新授权。"
    ])

    add_heading(doc, "四、蓝牙与硬件连接")
    add_bullets(doc, [
        "烧录完成后，拔掉 USB 线。",
        "在 Mac 蓝牙里搜索并连接 Vibe Coding Remote。",
        "如果搜不到，先在蓝牙里忽略旧的同名设备，再重新搜索。",
        "macOS 可能弹出键盘设置助理，可以关闭或跳过。"
    ])

    add_heading(doc, "五、一键烧录注意事项")
    add_bullets(doc, [
        "使用能传数据的 USB 线，不要只用充电线。",
        "烧录过程中不要拔线，不要让电脑睡眠。",
        "如果烧录失败，通常可以让 StickS3 重新进入下载模式后再烧一次。",
        "烧录成功日志里会出现 Hash of data verified。"
    ])

    add_heading(doc, "六、常见问题速查")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "现象"
    hdr[1].text = "处理方式"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    rows = [
        ("无法验证开发者", "右键 App -> 打开，或到隐私与安全性里点“仍要打开”。"),
        ("提示 App 已损坏", "执行 xattr 命令移除下载隔离标记，再右键打开。"),
        ("辅助功能开了但 App 检测不到", "删除旧授权条目，重新添加 /Applications 里的 App。"),
        ("找不到蓝牙设备", "忽略旧设备，重新搜索 Vibe Coding Remote；必要时重启设备。"),
        ("右侧键不打开目标软件", "确认 App 里目标软件选择正确，并已开启辅助功能权限。"),
        ("按住语音无反应", "确认语音输入设置为“按住 Fn”或已录制正确快捷键。"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_width(table, [2.0, 4.35])

    add_heading(doc, "七、给测试用户的短版说明")
    doc.add_paragraph("可以直接复制下面这段发给朋友：")
    add_command(
        doc,
        "打开 DMG 后先拖到应用程序；如果打不开，终端执行：\\n"
        'xattr -dr com.apple.quarantine "/Applications/Vibe Coding手持输入助手.app"\\n'
        "然后右键 App 点打开，并在系统设置 -> 隐私与安全性 -> 辅助功能里打开它。"
    )

    add_heading(doc, "八、作者信息")
    doc.add_paragraph("作者：智多星")
    doc.add_paragraph("个人微信：369076317")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
